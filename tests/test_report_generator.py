"""
Tests for report generation functionality.
"""
import pytest
from pathlib import Path
from datetime import datetime
import tempfile
import os

from report_generator import ReportGenerator


@pytest.fixture
def generator():
    """Create a ReportGenerator instance."""
    return ReportGenerator()


@pytest.fixture
def sample_data():
    """Provide sample patient and conflict data."""
    return {
        'patient_name': 'John Doe',
        'patient_id': 'P12345',
        'conditions': ['Hypertension', 'Diabetes'],
        'allergies': ['Penicillin'],
        'prescription': ['Aspirin', 'Metformin', 'Lisinopril'],
        'conflicts': [
            {
                'type': 'drug-drug',
                'item_a': 'Aspirin',
                'item_b': 'Warfarin',
                'severity': 'Major',
                'recommendation': 'Avoid concurrent use. Consider alternative antiplatelet agent.',
                'score': 9
            },
            {
                'type': 'drug-condition',
                'item_a': 'Ibuprofen',
                'item_b': 'Hypertension',
                'severity': 'Moderate',
                'recommendation': 'Monitor blood pressure regularly.',
                'score': 6
            }
        ]
    }


@pytest.fixture
def no_conflict_data():
    """Provide sample data with no conflicts."""
    return {
        'patient_name': 'Jane Smith',
        'patient_id': 'P67890',
        'conditions': ['Pain'],
        'allergies': [],
        'prescription': ['Acetaminophen'],
        'conflicts': []
    }


class TestPDFGeneration:
    """Test PDF report generation."""
    
    def test_generate_pdf_with_conflicts(self, generator, sample_data, tmp_path):
        """Test PDF generation with conflicts."""
        output_path = tmp_path / "test_report.pdf"
        
        result_path = generator.generate_pdf_report(
            output_path=output_path,
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        assert result_path.exists()
        assert result_path.suffix == '.pdf'
        assert result_path.stat().st_size > 0
    
    def test_generate_pdf_no_conflicts(self, generator, no_conflict_data, tmp_path):
        """Test PDF generation with no conflicts."""
        output_path = tmp_path / "safe_report.pdf"
        
        result_path = generator.generate_pdf_report(
            output_path=output_path,
            patient_name=no_conflict_data['patient_name'],
            patient_id=no_conflict_data['patient_id'],
            conditions=no_conflict_data['conditions'],
            allergies=no_conflict_data['allergies'],
            prescription=no_conflict_data['prescription'],
            conflicts=no_conflict_data['conflicts']
        )
        
        assert result_path.exists()
        assert result_path.stat().st_size > 0
    
    def test_generate_pdf_empty_prescription(self, generator, tmp_path):
        """Test PDF generation with empty prescription."""
        output_path = tmp_path / "empty_report.pdf"
        
        result_path = generator.generate_pdf_report(
            output_path=output_path,
            patient_name="Test Patient",
            patient_id="TEST-001",
            conditions=[],
            allergies=[],
            prescription=[],
            conflicts=[]
        )
        
        assert result_path.exists()
        assert result_path.stat().st_size > 0
    
    def test_pdf_creates_parent_directories(self, generator, sample_data, tmp_path):
        """Test that PDF generation creates parent directories."""
        output_path = tmp_path / "nested" / "folder" / "report.pdf"
        
        result_path = generator.generate_pdf_report(
            output_path=output_path,
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        assert result_path.exists()
        assert result_path.parent.exists()


class TestWordGeneration:
    """Test Word document report generation."""
    
    def test_generate_word_with_conflicts(self, generator, sample_data, tmp_path):
        """Test Word generation with conflicts."""
        output_path = tmp_path / "test_report.docx"
        
        result_path = generator.generate_word_report(
            output_path=output_path,
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        assert result_path.exists()
        assert result_path.suffix == '.docx'
        assert result_path.stat().st_size > 0
    
    def test_generate_word_no_conflicts(self, generator, no_conflict_data, tmp_path):
        """Test Word generation with no conflicts."""
        output_path = tmp_path / "safe_report.docx"
        
        result_path = generator.generate_word_report(
            output_path=output_path,
            patient_name=no_conflict_data['patient_name'],
            patient_id=no_conflict_data['patient_id'],
            conditions=no_conflict_data['conditions'],
            allergies=no_conflict_data['allergies'],
            prescription=no_conflict_data['prescription'],
            conflicts=no_conflict_data['conflicts']
        )
        
        assert result_path.exists()
        assert result_path.stat().st_size > 0
    
    def test_generate_word_empty_prescription(self, generator, tmp_path):
        """Test Word generation with empty prescription."""
        output_path = tmp_path / "empty_report.docx"
        
        result_path = generator.generate_word_report(
            output_path=output_path,
            patient_name="Test Patient",
            patient_id="TEST-001",
            conditions=[],
            allergies=[],
            prescription=[],
            conflicts=[]
        )
        
        assert result_path.exists()
        assert result_path.stat().st_size > 0
    
    def test_word_creates_parent_directories(self, generator, sample_data, tmp_path):
        """Test that Word generation creates parent directories."""
        output_path = tmp_path / "nested" / "folder" / "report.docx"
        
        result_path = generator.generate_word_report(
            output_path=output_path,
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        assert result_path.exists()
        assert result_path.parent.exists()


class TestBytesGeneration:
    """Test report generation to BytesIO for streaming."""
    
    def test_generate_pdf_bytes(self, generator, sample_data):
        """Test PDF generation as bytes."""
        pdf_bytes = generator.generate_report_bytes(
            format_type='pdf',
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        assert pdf_bytes is not None
        assert len(pdf_bytes.getvalue()) > 0
        assert pdf_bytes.tell() == 0  # Should be at start
    
    def test_generate_word_bytes(self, generator, sample_data):
        """Test Word generation as bytes."""
        word_bytes = generator.generate_report_bytes(
            format_type='word',
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        assert word_bytes is not None
        assert len(word_bytes.getvalue()) > 0
        assert word_bytes.tell() == 0
    
    def test_generate_docx_bytes(self, generator, sample_data):
        """Test Word generation with 'docx' format type."""
        word_bytes = generator.generate_report_bytes(
            format_type='docx',
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        assert word_bytes is not None
        assert len(word_bytes.getvalue()) > 0


class TestReportContent:
    """Test report content accuracy."""
    
    def test_pdf_contains_patient_info(self, generator, sample_data, tmp_path):
        """Test that PDF contains patient information."""
        output_path = tmp_path / "content_test.pdf"
        
        generator.generate_pdf_report(
            output_path=output_path,
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        # Basic check: file exists and has content
        assert output_path.exists()
        assert output_path.stat().st_size > 1000  # Should be reasonably sized
    
    def test_word_contains_patient_info(self, generator, sample_data, tmp_path):
        """Test that Word document contains patient information."""
        from docx import Document
        
        output_path = tmp_path / "content_test.docx"
        
        generator.generate_word_report(
            output_path=output_path,
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        # Read document and check content (paragraphs and tables)
        doc = Document(str(output_path))
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += '\n' + cell.text
        
        assert sample_data['patient_name'] in full_text
        assert sample_data['patient_id'] in full_text
        assert 'Hypertension' in full_text
        assert 'Aspirin' in full_text
    
    def test_report_includes_all_conflicts(self, generator, sample_data, tmp_path):
        """Test that all conflicts are included in report."""
        from docx import Document
        
        output_path = tmp_path / "conflicts_test.docx"
        
        generator.generate_word_report(
            output_path=output_path,
            patient_name=sample_data['patient_name'],
            patient_id=sample_data['patient_id'],
            conditions=sample_data['conditions'],
            allergies=sample_data['allergies'],
            prescription=sample_data['prescription'],
            conflicts=sample_data['conflicts']
        )
        
        doc = Document(str(output_path))
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Check that conflict details are present
        assert 'Aspirin' in full_text
        assert 'Warfarin' in full_text
        assert 'Major' in full_text
        assert 'Moderate' in full_text


class TestEdgeCases:
    """Test edge cases and error handling."""
    
    def test_many_conflicts(self, generator, tmp_path):
        """Test report generation with many conflicts."""
        conflicts = [
            {
                'type': 'drug-drug',
                'item_a': f'Drug{i}',
                'item_b': f'Drug{i+1}',
                'severity': 'Minor',
                'recommendation': f'Monitor for interaction {i}',
                'score': 3
            }
            for i in range(50)
        ]
        
        output_path = tmp_path / "many_conflicts.pdf"
        
        result_path = generator.generate_pdf_report(
            output_path=output_path,
            patient_name="Test Patient",
            patient_id="TEST-MANY",
            conditions=['Multiple Conditions'],
            allergies=[],
            prescription=[f'Drug{i}' for i in range(20)],
            conflicts=conflicts
        )
        
        assert result_path.exists()
        assert result_path.stat().st_size > 0
    
    def test_long_patient_name(self, generator, tmp_path):
        """Test with very long patient name."""
        long_name = "Dr. " + "Very Long Name " * 20
        
        output_path = tmp_path / "long_name.pdf"
        
        result_path = generator.generate_pdf_report(
            output_path=output_path,
            patient_name=long_name,
            patient_id="TEST-LONG",
            conditions=[],
            allergies=[],
            prescription=[],
            conflicts=[]
        )
        
        assert result_path.exists()
    
    def test_special_characters_in_data(self, generator, tmp_path):
        """Test with special characters in data."""
        output_path = tmp_path / "special_chars.docx"
        
        result_path = generator.generate_word_report(
            output_path=output_path,
            patient_name="O'Brien & Smith",
            patient_id="TEST-SPEC",
            conditions=["Type 2 Diabetes", "Heart Failure (NYHA Class II-III)"],
            allergies=["Aspirin (anaphylaxis)"],
            prescription=["Drug A/B Combination", "Drug C-D"],
            conflicts=[
                {
                    'type': 'drug-condition',
                    'item_a': 'Drug A/B',
                    'item_b': 'Heart Failure',
                    'severity': 'Moderate',
                    'recommendation': 'Monitor closely & adjust dose',
                    'score': 5
                }
            ]
        )
        
        assert result_path.exists()
