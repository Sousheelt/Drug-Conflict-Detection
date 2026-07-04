"""
Report generation module for Drug Conflict Detection System.

Generates professional PDF and Word reports with:
- Patient information and prescription details
- Conflict analysis with severity breakdown
- Recommendations and risk assessment
- Visual charts and formatted tables
"""

from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
from io import BytesIO

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus import Image as RLImage
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:
    """Generate PDF and Word reports for drug conflict analysis."""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles for PDF generation."""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        # Section header style
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=12
        ))
        
        # Warning style
        self.styles.add(ParagraphStyle(
            name='Warning',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.red,
            leftIndent=20
        ))
    
    def generate_pdf_report(
        self,
        output_path: Path | str,
        patient_name: str,
        patient_id: str,
        conditions: List[str],
        allergies: List[str],
        prescription: List[str],
        conflicts: List[Dict[str, Any]],
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """Generate a PDF report of conflict analysis.
        
        Args:
            output_path: Path to save the PDF report
            patient_name: Patient's name
            patient_id: Patient identifier
            conditions: List of medical conditions
            allergies: List of allergies
            prescription: List of prescribed drugs
            conflicts: List of detected conflicts
            metadata: Optional additional metadata
            
        Returns:
            Path to generated PDF file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for document elements
        story = []
        
        # Title
        story.append(Paragraph("Drug Conflict Detection Report", self.styles['CustomTitle']))
        story.append(Spacer(1, 12))
        
        # Generation timestamp
        timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
        story.append(Paragraph(f"Generated: {timestamp}", self.styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Patient Information Section
        story.append(Paragraph("Patient Information", self.styles['SectionHeader']))
        patient_data = [
            ['Patient Name:', patient_name],
            ['Patient ID:', patient_id],
            ['Conditions:', ', '.join(conditions) if conditions else 'None'],
            ['Allergies:', ', '.join(allergies) if allergies else 'None']
        ]
        
        patient_table = Table(patient_data, colWidths=[2*inch, 4*inch])
        patient_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8e8e8')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(patient_table)
        story.append(Spacer(1, 20))
        
        # Prescription Section
        story.append(Paragraph("Prescribed Medications", self.styles['SectionHeader']))
        if prescription:
            for i, drug in enumerate(prescription, 1):
                story.append(Paragraph(f"{i}. {drug}", self.styles['Normal']))
        else:
            story.append(Paragraph("No medications prescribed", self.styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Conflict Analysis Section
        story.append(Paragraph("Conflict Analysis", self.styles['SectionHeader']))
        
        if conflicts:
            # Summary
            major = sum(1 for c in conflicts if c.get('severity') == 'Major')
            moderate = sum(1 for c in conflicts if c.get('severity') == 'Moderate')
            minor = sum(1 for c in conflicts if c.get('severity') == 'Minor')
            
            summary_text = f"<b>Total Conflicts:</b> {len(conflicts)} "
            summary_text += f"(<font color='red'>Major: {major}</font>, "
            summary_text += f"<font color='orange'>Moderate: {moderate}</font>, "
            summary_text += f"<font color='#DAA520'>Minor: {minor}</font>)"
            story.append(Paragraph(summary_text, self.styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Detailed conflicts
            for i, conflict in enumerate(conflicts, 1):
                severity = conflict.get('severity', 'Unknown')
                item_a = conflict.get('item_a', '')
                item_b = conflict.get('item_b', '')
                conflict_type = conflict.get('type', 'unknown')
                recommendation = conflict.get('recommendation', 'No recommendation provided')
                
                # Severity color
                color_map = {
                    'Major': 'red',
                    'Moderate': 'orange',
                    'Minor': '#DAA520'
                }
                color = color_map.get(severity, 'black')
                
                conflict_text = f"<b>{i}. <font color='{color}'>{severity}</font>:</b> "
                conflict_text += f"{item_a} + {item_b}"
                story.append(Paragraph(conflict_text, self.styles['Normal']))
                
                story.append(Paragraph(f"<i>Type:</i> {conflict_type}", self.styles['Normal']))
                story.append(Paragraph(f"<i>Recommendation:</i> {recommendation}", self.styles['Normal']))
                story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(
                "<font color='green'><b>No conflicts detected.</b> The prescribed medications "
                "appear to be safe based on the current knowledge base.</font>",
                self.styles['Normal']
            ))
        
        story.append(Spacer(1, 20))
        
        # Risk Assessment
        story.append(Paragraph("Risk Assessment", self.styles['SectionHeader']))
        if conflicts:
            if major > 0:
                risk_level = "HIGH RISK"
                risk_color = "red"
                risk_text = "Immediate review recommended. Major drug interactions detected."
            elif moderate > 0:
                risk_level = "MODERATE RISK"
                risk_color = "orange"
                risk_text = "Caution advised. Monitor patient closely for adverse effects."
            else:
                risk_level = "LOW RISK"
                risk_color = "#DAA520"
                risk_text = "Minor interactions detected. Standard monitoring recommended."
        else:
            risk_level = "MINIMAL RISK"
            risk_color = "green"
            risk_text = "No significant interactions detected based on current data."
        
        risk_para = f"<b><font color='{risk_color}'>{risk_level}</font></b>: {risk_text}"
        story.append(Paragraph(risk_para, self.styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Disclaimer
        story.append(Paragraph("Important Disclaimer", self.styles['SectionHeader']))
        disclaimer = (
            "This report is generated by an automated drug conflict detection system "
            "for educational and informational purposes only. It should NOT be used as "
            "the sole basis for clinical decision-making. Always consult with qualified "
            "healthcare professionals and refer to current clinical guidelines."
        )
        story.append(Paragraph(disclaimer, self.styles['Normal']))
        
        # Build PDF
        doc.build(story)
        return output_path
    
    def generate_word_report(
        self,
        output_path: Path | str,
        patient_name: str,
        patient_id: str,
        conditions: List[str],
        allergies: List[str],
        prescription: List[str],
        conflicts: List[Dict[str, Any]],
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """Generate a Word document report of conflict analysis.
        
        Args:
            output_path: Path to save the Word document
            patient_name: Patient's name
            patient_id: Patient identifier
            conditions: List of medical conditions
            allergies: List of allergies
            prescription: List of prescribed drugs
            conflicts: List of detected conflicts
            metadata: Optional additional metadata
            
        Returns:
            Path to generated Word document
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Drug Conflict Detection Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
        timestamp_para = doc.add_paragraph(f"Generated: {timestamp}")
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Patient Information
        doc.add_heading('Patient Information', 1)
        patient_table = doc.add_table(rows=4, cols=2)
        patient_table.style = 'Light Grid Accent 1'
        
        patient_table.rows[0].cells[0].text = 'Patient Name:'
        patient_table.rows[0].cells[1].text = patient_name
        patient_table.rows[1].cells[0].text = 'Patient ID:'
        patient_table.rows[1].cells[1].text = patient_id
        patient_table.rows[2].cells[0].text = 'Conditions:'
        patient_table.rows[2].cells[1].text = ', '.join(conditions) if conditions else 'None'
        patient_table.rows[3].cells[0].text = 'Allergies:'
        patient_table.rows[3].cells[1].text = ', '.join(allergies) if allergies else 'None'
        
        # Bold first column
        for row in patient_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Prescription
        doc.add_heading('Prescribed Medications', 1)
        if prescription:
            for i, drug in enumerate(prescription, 1):
                doc.add_paragraph(f"{i}. {drug}", style='List Number')
        else:
            doc.add_paragraph("No medications prescribed")
        
        doc.add_paragraph()
        
        # Conflict Analysis
        doc.add_heading('Conflict Analysis', 1)
        
        if conflicts:
            # Summary
            major = sum(1 for c in conflicts if c.get('severity') == 'Major')
            moderate = sum(1 for c in conflicts if c.get('severity') == 'Moderate')
            minor = sum(1 for c in conflicts if c.get('severity') == 'Minor')
            
            summary = doc.add_paragraph()
            summary.add_run(f"Total Conflicts: {len(conflicts)} ").bold = True
            summary.add_run(f"(Major: {major}, Moderate: {moderate}, Minor: {minor})")
            
            doc.add_paragraph()
            
            # Detailed conflicts
            for i, conflict in enumerate(conflicts, 1):
                severity = conflict.get('severity', 'Unknown')
                item_a = conflict.get('item_a', '')
                item_b = conflict.get('item_b', '')
                conflict_type = conflict.get('type', 'unknown')
                recommendation = conflict.get('recommendation', 'No recommendation provided')
                
                # Conflict heading
                conflict_para = doc.add_paragraph()
                conflict_run = conflict_para.add_run(f"{i}. {severity}: ")
                conflict_run.bold = True
                
                # Color by severity
                color_map = {
                    'Major': RGBColor(255, 0, 0),
                    'Moderate': RGBColor(255, 165, 0),
                    'Minor': RGBColor(218, 165, 32)
                }
                if severity in color_map:
                    conflict_run.font.color.rgb = color_map[severity]
                
                conflict_para.add_run(f"{item_a} + {item_b}")
                
                # Details
                type_para = doc.add_paragraph(style='List Bullet')
                type_para.add_run('Type: ').italic = True
                type_para.add_run(conflict_type)
                
                rec_para = doc.add_paragraph(style='List Bullet')
                rec_para.add_run('Recommendation: ').italic = True
                rec_para.add_run(recommendation)
                
                doc.add_paragraph()
        else:
            no_conflict = doc.add_paragraph("No conflicts detected. ")
            no_conflict_run = no_conflict.runs[0]
            no_conflict_run.bold = True
            no_conflict_run.font.color.rgb = RGBColor(0, 128, 0)
            doc.add_paragraph(
                "The prescribed medications appear to be safe based on the current knowledge base."
            )
        
        # Risk Assessment
        doc.add_heading('Risk Assessment', 1)
        if conflicts:
            if major > 0:
                risk_level = "HIGH RISK"
                risk_color = RGBColor(255, 0, 0)
                risk_text = "Immediate review recommended. Major drug interactions detected."
            elif moderate > 0:
                risk_level = "MODERATE RISK"
                risk_color = RGBColor(255, 165, 0)
                risk_text = "Caution advised. Monitor patient closely for adverse effects."
            else:
                risk_level = "LOW RISK"
                risk_color = RGBColor(218, 165, 32)
                risk_text = "Minor interactions detected. Standard monitoring recommended."
        else:
            risk_level = "MINIMAL RISK"
            risk_color = RGBColor(0, 128, 0)
            risk_text = "No significant interactions detected based on current data."
        
        risk_para = doc.add_paragraph()
        risk_run = risk_para.add_run(risk_level + ": ")
        risk_run.bold = True
        risk_run.font.color.rgb = risk_color
        risk_para.add_run(risk_text)
        
        doc.add_paragraph()
        
        # Disclaimer
        doc.add_heading('Important Disclaimer', 1)
        disclaimer = doc.add_paragraph(
            "This report is generated by an automated drug conflict detection system "
            "for educational and informational purposes only. It should NOT be used as "
            "the sole basis for clinical decision-making. Always consult with qualified "
            "healthcare professionals and refer to current clinical guidelines."
        )
        disclaimer.runs[0].italic = True
        
        # Save document
        doc.save(str(output_path))
        return output_path
    
    def generate_report_bytes(
        self,
        format_type: str,
        patient_name: str,
        patient_id: str,
        conditions: List[str],
        allergies: List[str],
        prescription: List[str],
        conflicts: List[Dict[str, Any]],
        metadata: Optional[Dict[str, Any]] = None
    ) -> BytesIO:
        """Generate report as bytes for streaming/download.
        
        Args:
            format_type: 'pdf' or 'word'
            Other args: Same as generate_pdf_report/generate_word_report
            
        Returns:
            BytesIO object containing the report
        """
        buffer = BytesIO()
        
        if format_type.lower() == 'pdf':
            # Create temporary path
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp_path = Path(tmp.name)
            
            # Generate PDF
            self.generate_pdf_report(
                tmp_path, patient_name, patient_id, conditions,
                allergies, prescription, conflicts, metadata
            )
            
            # Read into buffer
            with open(tmp_path, 'rb') as f:
                buffer.write(f.read())
            
            # Cleanup
            tmp_path.unlink()
            
        elif format_type.lower() in ['word', 'docx']:
            # Create temporary path
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = Path(tmp.name)
            
            # Generate Word document using the full method
            self.generate_word_report(
                tmp_path, patient_name, patient_id, conditions,
                allergies, prescription, conflicts, metadata
            )
            
            # Read into buffer
            with open(tmp_path, 'rb') as f:
                buffer.write(f.read())
            
            # Cleanup
            tmp_path.unlink()
        
        buffer.seek(0)
        return buffer
